"""Servicio de generacion de descripcion tecnica del predio.

Genera el texto descriptivo (linderos tecnicos) y permite exportarlo
a formato DOCX.
"""

from __future__ import annotations

import io
from typing import Any

from app.services.linderos_svc import obtener_secuencia_puntos, calcular_distancia_total
from app.services.coordenadas_svc import calcular_rumbo_por_coords
from app.core.geometry import build_indexes


# ---------------------------------------------------------------------------
# Generacion de descripcion textual
# ---------------------------------------------------------------------------

def generar_descripcion_textual(project_data: dict[str, Any]) -> str:
    """Genera la descripcion tecnica completa del predio en texto plano.

    Utiliza los datos del proyecto (datos_predio, segmentos, linderos,
    rumbos_asignados) para construir la descripcion de los linderos
    tecnicos siguiendo el formato topografico colombiano.

    Parametros:
        project_data: diccionario con los datos completos del proyecto
            (equivalente a ProjectResponse serializado).

    Retorna:
        Cadena de texto con la descripcion tecnica completa.
    """
    datos_predio = project_data.get("datos_predio", {})
    segmentos = project_data.get("segmentos", [])
    linderos = project_data.get("linderos", {})
    rumbos_asignados = project_data.get("rumbos_asignados", {
        "NORTE": [], "ESTE": [], "SUR": [], "OESTE": []
    })

    nombre_predio = datos_predio.get("nombre_predio", "SIN NOMBRE")
    municipio = datos_predio.get("municipio", "")
    departamento = datos_predio.get("departamento", "")
    vereda = datos_predio.get("vereda", "")

    area_m2 = project_data.get("area_m2")
    perimetro_m = project_data.get("perimetro_m")

    lines: list[str] = []

    # Encabezado
    lines.append("DESCRIPCION TECNICA DE LINDEROS")
    lines.append("=" * 50)
    lines.append("")
    lines.append(f"Predio: {nombre_predio}")
    if municipio:
        lines.append(f"Municipio: {municipio}")
    if departamento:
        lines.append(f"Departamento: {departamento}")
    if vereda:
        lines.append(f"Vereda: {vereda}")
    lines.append("")

    if area_m2 is not None:
        lines.append(f"Area: {area_m2:.2f} m2")
    if perimetro_m is not None:
        lines.append(f"Perimetro: {perimetro_m:.2f} m")
    lines.append("")

    # Descripcion de linderos por rumbo
    indexes = build_indexes(segmentos)
    orden_rumbos = ["NORTE", "ESTE", "SUR", "OESTE"]

    for rumbo in orden_rumbos:
        indices_linderos = rumbos_asignados.get(rumbo, [])
        if not indices_linderos:
            continue

        lines.append(f"--- {rumbo} ---")
        lines.append("")

        for idx in indices_linderos:
            lindero = linderos.get(str(idx), linderos.get(idx, {}))
            if not lindero:
                continue

            # Obtener datos del lindero
            if isinstance(lindero, dict):
                inicial = lindero.get("inicial", "")
                final = lindero.get("final", "")
                colindante = lindero.get("colindante", "")
                descripcion_manual = lindero.get("descripcion", "")
            else:
                continue

            # Calcular secuencia y distancia
            secuencia = obtener_secuencia_puntos(segmentos, inicial, final)
            distancia = calcular_distancia_total(segmentos, secuencia)

            # Construir descripcion del tramo
            if secuencia and len(secuencia) >= 2:
                puntos_str = " - ".join(secuencia)
                tramo_desc = (
                    f"Del punto {inicial} al punto {final}, "
                    f"pasando por los puntos {puntos_str}, "
                    f"en una distancia de {distancia:.2f} metros"
                )
            else:
                tramo_desc = f"Del punto {inicial} al punto {final}"

            if colindante:
                tramo_desc += f", colindando con {colindante}"

            tramo_desc += "."

            if descripcion_manual:
                tramo_desc = descripcion_manual

            lines.append(tramo_desc)
            lines.append("")

    # Cierre
    lines.append("=" * 50)
    lines.append("Fin de la descripcion tecnica.")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Generacion de documento DOCX
# ---------------------------------------------------------------------------

def generar_docx(project_data: dict[str, Any]) -> bytes:
    """Genera un archivo DOCX con la descripcion tecnica del predio.

    Retorna los bytes del archivo DOCX listo para descargar.

    NOTA: Implementacion basica usando python-docx. Se puede ampliar
    con formatos mas elaborados (tablas, logos, etc.).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback: si python-docx no esta instalado, retornar
        # un archivo de texto plano como .docx (no ideal pero funcional)
        texto = generar_descripcion_textual(project_data)
        return texto.encode("utf-8")

    doc = Document()

    # Configurar estilos basicos
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    datos_predio = project_data.get("datos_predio", {})
    nombre_predio = datos_predio.get("nombre_predio", "SIN NOMBRE")

    # Titulo
    titulo = doc.add_heading("DESCRIPCION TECNICA DE LINDEROS", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos del predio
    doc.add_paragraph("")
    doc.add_paragraph(f"Predio: {nombre_predio}")

    municipio = datos_predio.get("municipio", "")
    departamento = datos_predio.get("departamento", "")
    vereda = datos_predio.get("vereda", "")

    if municipio:
        doc.add_paragraph(f"Municipio: {municipio}")
    if departamento:
        doc.add_paragraph(f"Departamento: {departamento}")
    if vereda:
        doc.add_paragraph(f"Vereda: {vereda}")

    area_m2 = project_data.get("area_m2")
    perimetro_m = project_data.get("perimetro_m")

    if area_m2 is not None:
        doc.add_paragraph(f"Area: {area_m2:.2f} m2")
    if perimetro_m is not None:
        doc.add_paragraph(f"Perimetro: {perimetro_m:.2f} m")

    doc.add_paragraph("")

    # Descripcion tecnica
    texto = generar_descripcion_textual(project_data)
    # Saltar el encabezado (ya lo pusimos con formato)
    lineas = texto.split("\n")
    en_cuerpo = False
    for linea in lineas:
        if linea.startswith("---"):
            doc.add_heading(linea.replace("---", "").strip(), level=2)
            en_cuerpo = True
            continue
        if en_cuerpo and linea.strip():
            doc.add_paragraph(linea.strip())

    # Guardar a bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
